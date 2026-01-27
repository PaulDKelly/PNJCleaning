import docx
import os

def extract_docx_info(docx_path):
    doc = docx.Document(docx_path)
    output = []
    
    output.append(f"--- TEXT FROM {os.path.basename(docx_path)} ---")
    for para in doc.paragraphs:
        if para.text.strip():
            output.append(para.text)
            
    output.append("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        output.append(f"\nTable {i+1}:")
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            output.append(" | ".join(row_data))
            
    return "\n".join(output)

if __name__ == "__main__":
    path = r"E:\Code\Projects\PNJCleaning\08012026 Chuck and Blade Dolphin Centre.docx"
    content = extract_docx_info(path)
    with open("docx_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Done writing to docx_content.txt")
